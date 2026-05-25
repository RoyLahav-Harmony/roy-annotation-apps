import io
import streamlit as st
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor

st.set_page_config(page_title="Agent Script Summary", layout="wide")
st.title("Agent Script Summary")

# ── Text input ─────────────────────────────────────────────────────────────────
pasted = st.text_area("Paste agent script JSON here", height=250, placeholder='{ "goals": [ ... ] }')

if not pasted.strip():
    st.info("Paste an agent script JSON above to begin.")
    st.stop()


# ── Helpers ────────────────────────────────────────────────────────────────────
def safe_str(value, default=""):
    """Return value as a stripped string, or default if it's not a string."""
    return value.strip() if isinstance(value, str) else default


def extract_messages(field):
    """Return a flat list of non-empty utterance strings from any messages format."""
    if not field:
        return []
    if isinstance(field, str):
        return [field.strip()] if field.strip() else []
    if isinstance(field, list):
        out = []
        for m in field:
            if isinstance(m, str) and m.strip():
                out.append(m.strip())
            elif isinstance(m, dict):
                text = safe_str(m.get("message", ""))
                if text:
                    out.append(text)
        return out
    return []


def walk_goals(goals, parent_path=""):
    """
    Recursively yield (label, [utterances]) for every extractable piece
    of a goal tree, including refusal handling, choice acknowledges,
    copy_from references, and nested sub-goals.
    """
    for goal in (goals or []):
        if goal.get("enabled") is False:
            continue

        name = safe_str(goal.get("name", ""))
        path = parent_path if goal.get("appended") and parent_path else (
            f"{parent_path} / {name}" if parent_path else name
        )

        utterances = extract_messages(goal.get("messages"))
        if not utterances and goal.get("copy_from"):
            utterances = [f"[external template: {goal['copy_from']}]"]
        if utterances:
            yield path, utterances

        refusal = extract_messages(goal.get("refusal_handling"))
        if refusal:
            yield f"{path} (refusal handling)", refusal

        yield from walk_goals(goal.get("goals", []), path)

        for choice in (goal.get("choices") or []):
            choice_name = safe_str(choice.get("name", ""))
            ack = safe_str(choice.get("acknowledge", ""))
            if ack:
                yield f"{path} = {choice_name} acknowledge", [ack]
            choice_path = f"{path} / {choice_name}" if choice_name else path
            yield from walk_goals(choice.get("goals", []), choice_path)


def walk_experiments(experiments):
    """Yield (label, [utterances]) from experiment test-group goal overrides."""
    for exp in (experiments or []):
        for group in (exp.get("test_groups") or []):
            goals_update = group.get("goals_update") or {}
            for goal_name, overrides in goals_update.items():
                msgs = extract_messages(overrides.get("messages"))
                if msgs:
                    yield f"{goal_name} (experiment: {group.get('name', exp.get('name', 'variant'))})", msgs


def walk_metadata(default_metadata):
    """
    Yield (label, [utterance]) for default_metadata values that look like
    natural-language utterances (strings > 20 chars, not pure variable refs).
    """
    for key, value in (default_metadata or {}).items():
        if not isinstance(value, str):
            continue
        text = value.strip()
        if len(text) < 20:
            continue
        if text.startswith("http") or (text.startswith("{%") and text.endswith("%}")):
            continue
        label = key.strip("_%").replace("_", " ")
        yield f"default: {label}", [text]


# ── Main extraction ────────────────────────────────────────────────────────────
@st.cache_data
def build_summary(text: str):
    script = json.loads(text)
    rows = []

    def add(label, utterances):
        unique = list(dict.fromkeys(utterances))
        rows.append({
            "Goal": label,
            "Utterances": " | ".join(unique),
            "# Variants": len(unique),
        })

    for label, utts in walk_goals(script.get("goals", [])):
        add(label, utts)

    for label, utts in walk_experiments(script.get("experiments", [])):
        add(label, utts)

    for label, utts in walk_metadata(script.get("default_metadata")):
        add(label, utts)

    return pd.DataFrame(rows, columns=["Goal", "Utterances", "# Variants"])


try:
    df = build_summary(pasted)
except json.JSONDecodeError as e:
    st.error(f"Invalid JSON: {e}")
    st.stop()
except Exception as e:
    st.error(f"Could not parse script: {e}")
    st.stop()

if df.empty:
    st.warning("No goals with utterances found in this script.")
    st.stop()

# ── Metrics ────────────────────────────────────────────────────────────────────
c1, c2, c3 = st.columns(3)
c1.metric("Goals found", len(df))
c2.metric("Total utterance variants", int(df["# Variants"].sum()))
c3.metric("Goals with multiple variants", int((df["# Variants"] > 1).sum()))

st.markdown("---")

# ── Filters ────────────────────────────────────────────────────────────────────
search = st.text_input("Filter by goal name or utterance text")
only_multi = st.checkbox("Show only goals with multiple utterance variants")

filtered = df.copy()
if search:
    q = search.lower()
    filtered = filtered[
        filtered["Goal"].str.lower().str.contains(q, na=False)
        | filtered["Utterances"].str.lower().str.contains(q, na=False)
    ]
if only_multi:
    filtered = filtered[filtered["# Variants"] > 1]

st.caption(f"Showing {len(filtered)} of {len(df)} goals")

# ── Table ──────────────────────────────────────────────────────────────────────
st.dataframe(
    filtered,
    use_container_width=True,
    hide_index=True,
    column_config={
        "Goal":       st.column_config.TextColumn("Goal", width="medium"),
        "Utterances": st.column_config.TextColumn("Utterances (| separated)", width="large"),
        "# Variants": st.column_config.NumberColumn("# Variants", width="small"),
    },
)

# ── Word doc export ────────────────────────────────────────────────────────────
def build_docx(data: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading("Agent Script Summary", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Goal", "Utterances", "# Variants"]):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(
            __import__("lxml.etree", fromlist=["etree"]).etree.fromstring(
                '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' w:val="clear" w:color="auto" w:fill="2E5FA3"/>'
            )
        )

    # Data rows
    for _, row in data.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Goal"])
        cells[1].text = str(row["Utterances"])
        cells[2].text = str(row["# Variants"])
        for cell in cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


st.markdown("---")
st.download_button(
    label="⬇️ Download as Word doc",
    data=build_docx(filtered),
    file_name="agent_script_summary.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
