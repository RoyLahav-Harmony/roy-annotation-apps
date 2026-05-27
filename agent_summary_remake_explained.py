# ══════════════════════════════════════════════════════════════════════════════
# agent_summary_remake_explained.py
# Annotated version of agent_summary_remake.py for documentation purposes.
#
# Purpose: Takes a raw agent script JSON (the configuration file that defines
# what an AI voice agent says and does) and renders it as a readable table
# showing every goal, its utterance variants, and optionally exports to Word.
# ══════════════════════════════════════════════════════════════════════════════

import io
import streamlit as st
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor

st.set_page_config(page_title="Agent Script Summary", layout="wide")
st.title("Agent Script Summary")


# ── Text input ─────────────────────────────────────────────────────────────────
# The user pastes the raw JSON of an agent script directly into this text box.
# (There's no file upload here — the JSON is pasted as text.)

pasted = st.text_area("Paste agent script JSON here", height=250, placeholder='{ "goals": [ ... ] }')

if not pasted.strip():
    st.info("Paste an agent script JSON above to begin.")
    st.stop()  # Nothing below runs until there's some input.


# ── Helper: safe string extraction ─────────────────────────────────────────────
def safe_str(value, default=""):
    """Return value as a stripped string, or default if it's not a string."""
    return value.strip() if isinstance(value, str) else default


# ── Helper: message extraction ─────────────────────────────────────────────────
def extract_messages(field):
    """
    Agent scripts store utterances in several possible formats:
      - A plain string: "Hello, how can I help?"
      - A list of strings: ["Hello!", "Hi there!"]
      - A list of dicts: [{"message": "Hello!"}, ...]

    This function normalises all three formats into a flat list of strings.
    """
    if not field:
        return []
    if isinstance(field, str):
        return [field.strip()] if field.strip() else []
    if isinstance(field, list):
        out = []
        for m in field:
            if isinstance(m, str) and m.strip():
                out.append(m.strip())
            elif isinstance(m, dict):
                # Dict format — the text lives under the "message" key.
                text = safe_str(m.get("message", ""))
                if text:
                    out.append(text)
        return out
    return []


# ── Helper: recursive goal tree walker ─────────────────────────────────────────
def walk_goals(goals, parent_path=""):
    """
    Agent scripts are structured as a tree of "goals". Each goal can have:
      - messages: the utterances the agent says when reaching this goal.
      - refusal_handling: what the agent says if the user refuses/pushes back.
      - choices: branching options the user can select, each with sub-goals.
      - goals: nested child goals (recursive structure).
      - copy_from: a reference to an external template instead of inline messages.
      - appended: if True, this goal's name is not added to the path (it appends
                  to the parent path instead of creating a new level).
      - enabled: if False, the goal is disabled and should be skipped entirely.

    Yields (label, [utterances]) for every node that has displayable content.
    The label is a slash-separated path showing where in the tree this goal sits.
    """
    for goal in (goals or []):
        # Skip disabled goals — they're defined but not active in the script.
        if goal.get("enabled") is False:
            continue

        name = safe_str(goal.get("name", ""))
        # "appended" goals extend the parent path rather than adding a new level.
        path = parent_path if goal.get("appended") and parent_path else (
            f"{parent_path} / {name}" if parent_path else name
        )

        # Extract the direct utterances for this goal node.
        utterances = extract_messages(goal.get("messages"))
        # If there are no inline messages but a copy_from reference, note it.
        if not utterances and goal.get("copy_from"):
            utterances = [f"[external template: {goal['copy_from']}]"]
        if utterances:
            yield path, utterances

        # Refusal handling — what the agent says if the user objects or refuses.
        refusal = extract_messages(goal.get("refusal_handling"))
        if refusal:
            yield f"{path} (refusal handling)", refusal

        # Recurse into nested child goals.
        yield from walk_goals(goal.get("goals", []), path)

        # Process choice branches — each choice can have an acknowledgement
        # (what the agent says when that choice is selected) and its own sub-goals.
        for choice in (goal.get("choices") or []):
            choice_name = safe_str(choice.get("name", ""))
            ack = safe_str(choice.get("acknowledge", ""))
            if ack:
                yield f"{path} = {choice_name} acknowledge", [ack]
            choice_path = f"{path} / {choice_name}" if choice_name else path
            yield from walk_goals(choice.get("goals", []), choice_path)


# ── Helper: experiment variant walker ──────────────────────────────────────────
def walk_experiments(experiments):
    """
    Some agent scripts have A/B test experiments that override goal messages
    for specific test groups. This extracts those overrides alongside a label
    showing which experiment and variant they belong to.
    """
    for exp in (experiments or []):
        for group in (exp.get("test_groups") or []):
            goals_update = group.get("goals_update") or {}
            for goal_name, overrides in goals_update.items():
                msgs = extract_messages(overrides.get("messages"))
                if msgs:
                    yield f"{goal_name} (experiment: {group.get('name', exp.get('name', 'variant'))})", msgs


# ── Helper: default metadata walker ────────────────────────────────────────────
def walk_metadata(default_metadata):
    """
    default_metadata is a key-value store on the agent for template variables
    like company name, agent persona, etc. Some of these are short IDs (skipped),
    URLs (skipped), or template expressions (skipped). Long plain-language strings
    are likely natural-language utterances worth showing in the summary.
    """
    for key, value in (default_metadata or {}).items():
        if not isinstance(value, str):
            continue
        text = value.strip()
        # Skip very short values, URLs, and Jinja-style template expressions.
        if len(text) < 20:
            continue
        if text.startswith("http") or (text.startswith("{%") and text.endswith("%}")):
            continue
        # Convert the key name to a readable label (e.g. "_company_name_" → "company name").
        label = key.strip("_%").replace("_", " ")
        yield f"default: {label}", [text]


# ── Main extraction function ───────────────────────────────────────────────────
# @st.cache_data: only re-parses if the pasted text changes.

@st.cache_data
def build_summary(text: str):
    """
    Parse the agent script JSON and build a DataFrame with one row per goal node.
    Columns:
      - Goal: the slash-separated path identifying the node in the script tree.
      - Utterances: all utterance variants for this node, joined with " | ".
      - # Variants: how many distinct utterances this node has.
    """
    script = json.loads(text)
    rows = []

    def add(label, utterances):
        # Deduplicate utterances while preserving order.
        unique = list(dict.fromkeys(utterances))
        rows.append({
            "Goal":       label,
            "Utterances": " | ".join(unique),
            "# Variants": len(unique),
        })

    # Walk the three sources of content in the script.
    for label, utts in walk_goals(script.get("goals", [])):
        add(label, utts)

    for label, utts in walk_experiments(script.get("experiments", [])):
        add(label, utts)

    for label, utts in walk_metadata(script.get("default_metadata")):
        add(label, utts)

    return pd.DataFrame(rows, columns=["Goal", "Utterances", "# Variants"])


# ── Parse and validate ─────────────────────────────────────────────────────────
try:
    df = build_summary(pasted)
except json.JSONDecodeError as e:
    st.error(f"Invalid JSON: {e}")
    st.stop()
except Exception as e:
    st.error(f"Could not parse script: {e}")
    st.stop()

if df.empty:
    st.warning("No goals with utterances found in this script.")
    st.stop()


# ── Summary metrics ────────────────────────────────────────────────────────────
# Quick overview of the script's scale.
c1, c2, c3 = st.columns(3)
c1.metric("Goals found",                    len(df))
c2.metric("Total utterance variants",       int(df["# Variants"].sum()))
c3.metric("Goals with multiple variants",   int((df["# Variants"] > 1).sum()))

st.markdown("---")


# ── Filters ────────────────────────────────────────────────────────────────────
# Let the user narrow down the table by text search or by multi-variant goals only.

search     = st.text_input("Filter by goal name or utterance text")
only_multi = st.checkbox("Show only goals with multiple utterance variants")

filtered = df.copy()
if search:
    q = search.lower()
    filtered = filtered[
        filtered["Goal"].str.lower().str.contains(q, na=False)
        | filtered["Utterances"].str.lower().str.contains(q, na=False)
    ]
if only_multi:
    filtered = filtered[filtered["# Variants"] > 1]

st.caption(f"Showing {len(filtered)} of {len(df)} goals")


# ── Table display ──────────────────────────────────────────────────────────────
# The column_config controls column widths and display labels.
st.dataframe(
    filtered,
    use_container_width=True,
    hide_index=True,
    column_config={
        "Goal":       st.column_config.TextColumn("Goal",                        width="medium"),
        "Utterances": st.column_config.TextColumn("Utterances (| separated)",    width="large"),
        "# Variants": st.column_config.NumberColumn("# Variants",                width="small"),
    },
)


# ── Word document export ───────────────────────────────────────────────────────
# Builds a .docx file in memory using the python-docx library and offers it
# for download. The document includes a pre-filled header with blank fields
# for the reviewer to complete (company name, domain, etc.).

def build_docx(data: pd.DataFrame) -> bytes:
    from lxml import etree

    doc = Document()
    doc.add_heading("Agent Script Summary", level=1)

    # Blank fields at the top for the reviewer to fill in manually.
    for label in ["Company name:", "Location:", "Domain:", "Agent name:", "Overall description:", "Hand off:"]:
        doc.add_paragraph(label)

    # Create a 3-column table with a styled header row.
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Style the header row: bold white text on a blue background.
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Goal", "Utterances", "# Variants"]):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White text.
        # The blue background requires injecting raw OOXML because python-docx
        # doesn't expose cell shading through its high-level API.
        shd = etree.fromstring(
            '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:val="clear" w:color="auto" w:fill="2E5FA3"/>'
        )
        cell._tc.get_or_add_tcPr().append(shd)

    # Add one row per goal in the filtered DataFrame.
    for _, row in data.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Goal"])
        cells[1].text = str(row["Utterances"])
        cells[2].text = str(row["# Variants"])
        for cell in cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Save to an in-memory buffer and return the raw bytes for the download button.
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


st.markdown("---")
st.download_button(
    label="⬇️ Download as Word doc",
    data=build_docx(filtered),
    file_name="agent_script_summary.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
